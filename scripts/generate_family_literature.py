"""Generate concise, sales- and SEO-oriented product literature for every family.

Runs entirely locally, no LLM calls. Mines the existing deep-dive markdown docs,
families.json, and the SKU catalogue, then produces for each family:
  - output/literature/<manufacturer>/<slug>.md    (customer-facing page)
  - output/literature/<manufacturer>/<slug>.docx  (matching Word document)

The copy follows the structure of the Thermotec 4-Zero literature draft
(product description, key features, applications/selection checklist, range
table, technical data, compliance, installation, safety, sustainability,
warranty, spec clause, source register, review actions) while keeping the
repo's safety rules: only facts already present in the knowledge base are used,
and every document keeps the human-review gates and "draft pending TDS
confirmation" status.

Only regenerates files whose inputs changed (content hashing), so it is cheap
to run repeatedly in the background.

Usage:
    python scripts/generate_family_literature.py              # all families
    python scripts/generate_family_literature.py --only Autex # one manufacturer
    python scripts/generate_family_literature.py --dry-run    # report only
"""
from __future__ import annotations

import argparse
import hashlib
import json
import re
import sys
from pathlib import Path

import pandas as pd

ROOT = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(ROOT))

OUT_DIR = ROOT / "output" / "literature"
SKU_CSV = ROOT / "data" / "processed" / "product_catalogue_skus.csv"
STATE_FILE = OUT_DIR / ".literature_state.json"
GENERATOR_VERSION = "2"  # bump when the template changes so all outputs regenerate

CATEGORY_TAGLINES = {
    "Batt": "bulk insulation batts for thermal and acoustic performance",
    "Board": "rigid insulation boards for continuous thermal performance",
    "Reflective": "reflective foil insulation for radiant heat control",
    "Wrap": "reflective membrane for weather protection and condensation control",
    "Pipe": "pre-formed pipe insulation for thermal efficiency and condensation control",
    "Panel": "acoustic panels for sound absorption and interior finish",
    "Accessory": "installation accessories and fixings",
}
INTENT_TERMS = {
    "thermal": "thermal insulation", "acoustic": "acoustic insulation",
    "wall": "wall insulation", "ceiling": "ceiling insulation",
    "roof": "roof insulation", "floor": "floor insulation",
    "underfloor": "underfloor insulation", "pipe": "pipe insulation",
    "duct": "duct insulation", "shed": "shed insulation",
}


def clean(text: str) -> str:
    return re.sub(r"\s+", " ", str(text or "")).strip()


def rating_summary(values: pd.Series) -> str:
    """Collapse noisy per-SKU rating text (e.g. 'Review: R1.2 | R2.5') to a
    compact, de-duplicated summary for a document cell."""
    found: list[str] = []
    for value in values.dropna().astype(str):
        for token in re.findall(r"R\d+(?:\.\d+)?", value):
            if token not in found:
                found.append(token)
    return ", ".join(found[:4]) if found else "Not stated"


def slugify(name: str) -> str:
    slug = re.sub(r"[^a-zA-Z0-9]+", "_", name).strip("_").lower()
    return slug[:60]


def parse_front_matter(text: str) -> dict:
    match = re.match(r"^---\n(.*?)\n---\n", text, re.S)
    if not match:
        return {}
    data = {}
    for line in match.group(1).splitlines():
        if ":" in line and not line.startswith(" "):
            key, _, value = line.partition(":")
            data[key.strip()] = value.strip()
    return data


def section(text: str, heading: str) -> str:
    pattern = rf"(?ms)^## {re.escape(heading)}\n(.*?)(?=^## |\Z)"
    match = re.search(pattern, text)
    return match.group(1).strip() if match else ""


def extract_features(text: str) -> list[str]:
    match = re.search(r"features and benefits include:(.*?)(?:This family covers|\n\n|$)", text, re.S | re.I)
    if not match:
        return []
    raw = match.group(1).replace("\n", " ")
    parts = [clean(part).rstrip(".") for part in re.split(r";|\.\s+(?=[A-Z])", raw)]
    return [part for part in parts if len(part) > 8][:8]


def extract_description(text: str) -> str:
    body = section(text, "Canonical description")
    body = re.sub(r"Manufacturer-published features and benefits include:.*?(?=(?:This family covers|\n\n|$))", "", body, flags=re.S | re.I)
    body = re.sub(r"This family covers .*? SKU variant.*?\.", "", body)
    body = re.sub(r"It is not a complete compliant building system.*", "", body, flags=re.S)
    sentences = re.split(r"(?<=[.!?])\s+", clean(body))
    return " ".join(sentences[:3])


def extract_limitations(text: str) -> list[str]:
    body = section(text, "Manufacturer-stated limitations and warnings")
    bullets = re.findall(r"(?m)^- (.+)$", body)
    return [clean(b) for b in bullets if "No manufacturer limitations" not in b][:6]


def extract_grade_rows(text: str) -> list[dict]:
    body = section(text, "Grade and catalogue reconciliation")
    lines = [line for line in body.splitlines() if line.strip().startswith("|")]
    rows = []
    for line in lines:
        cells = [clean(cell) for cell in line.strip().strip("|").split("|")]
        if len(cells) == 5 and cells[0] not in ("Rating (as supplied)", "---"):
            rows.append({"rating": cells[0], "rating_type": cells[1], "thickness": cells[2], "dimensions": cells[3], "sku_count": cells[4]})
    return rows


def seo_keywords(name: str, manufacturer: str, category: str, applications: list[str], ratings: set[str]) -> list[str]:
    keywords = [name, f"{manufacturer} {category.lower()}" if category else manufacturer]
    lowered = " ".join([name, *applications]).casefold()
    for term, phrase in INTENT_TERMS.items():
        if term in lowered:
            keywords.append(phrase)
    if "R" in " ".join(ratings):
        keywords.append("R-value insulation")
    keywords += ["insulation Australia", f"{manufacturer} Australia"]
    seen, ordered = set(), []
    for keyword in keywords:
        key = keyword.casefold()
        if key not in seen:
            seen.add(key)
            ordered.append(keyword)
    return ordered[:10]


def title_case_manufacturer(manufacturer_dir: str) -> str:
    return manufacturer_dir.title()


def build_markdown(family: dict, fm: dict, text: str, skus: pd.DataFrame) -> tuple[str, dict]:
    name = family["name"]
    manufacturer = family["manufacturer"]
    category = clean(family.get("category", "")).replace(" insulation", "").replace(" Insulation", "")
    applications = family.get("applications", [])
    description = extract_description(text)
    features = extract_features(text)
    limitations = extract_limitations(text)
    grade_rows = extract_grade_rows(text)
    material = fm.get("material") or (clean(skus["material_type"].iloc[0]) if not skus.empty else "")
    tds_url = family.get("source_url") or fm.get("official_datasheet_url", "")
    sds_url = fm.get("official_sds_url", "")
    ratings: set[str] = set()
    if not skus.empty:
        for column in ("thermal_r_value", "acoustic_rw", "nrc_aw"):
            for value in skus[column].dropna().astype(str):
                for token in re.findall(r"(?:R|Rw|NRC)\s?\d+(?:\.\d+)?", value):
                    ratings.add(token)
    rating_text = ", ".join(sorted(ratings)[:8]) if ratings else "Not yet extracted per SKU"
    seo_description = clean(description) or f"{name} is a {category.lower()} insulation product family from {manufacturer}. View the catalogue range, applications and specification starting point."
    keywords = seo_keywords(name, manufacturer, category, applications, ratings)
    tagline = CATEGORY_TAGLINES.get(category, f"{category.lower()} insulation product")

    feature_md = "\n".join(f"- {feature}." for feature in features) or "- Refer to the manufacturer datasheet for published features."
    apps_md = "\n".join(f"- {app}" for app in applications) or "- General building applications."
    limits_md = "\n".join(f"- {item}" for item in limitations) or "- No manufacturer limitations extracted yet; treat all claims as unverified until the TDS/SDS is reviewed."
    kw_md = ", ".join(keywords)

    range_rows = ""
    if not skus.empty:
        range_rows = "| SKU | Product | Published rating |\n| --- | --- | --- |\n"
        for _, sku in skus.head(25).iterrows():
            rating = rating_summary(pd.Series([sku["thermal_r_value"], sku["acoustic_rw"], sku["nrc_aw"]]))
            range_rows += f"| {clean(sku['our_sku'])} | {clean(sku['product_name'])} | {rating} |\n"
        if len(skus) > 25:
            range_rows += f"\n_{len(skus) - 25} further catalogue variants not listed here._\n"
    elif grade_rows:
        range_rows = "| Rating | Type | Thickness | Dimensions | SKUs |\n| --- | --- | --- | --- | --- |\n"
        for row in grade_rows:
            range_rows += f"| {row['rating']} | {row['rating_type']} | {row['thickness']} | {row['dimensions']} | {row['sku_count']} |\n"
    else:
        range_rows = "_Range not yet extracted; confirm variants against the current manufacturer TDS._\n"

    meta = {
        "name": name, "manufacturer": manufacturer, "category": category,
        "tagline": tagline, "keywords": keywords, "description": description,
        "features": features, "limitations": limitations, "material": material,
        "tds_url": tds_url, "sds_url": sds_url, "applications": applications,
        "skus": skus, "grade_rows": grade_rows,
    }

    md = f"""---
title: "{name} - {category} Insulation | {manufacturer}"
description: "{seo_description[:150]}"
keywords: "{kw_md}"
status: "Draft - pending manufacturer TDS/SDS confirmation"
family_id: {family['family_id']}
---

# {name}

**{manufacturer} {category}** — {tagline}.

{description or f"{name} is a {category.lower()} product family from {manufacturer}. Confirm the current published specification against the manufacturer datasheet before quoting."}

## Key features

{feature_md}

## Applications and selection

{apps_md}

**Selection checklist**

1. Confirm the application (wall, ceiling, floor, roof, pipe or service) matches the family.
2. Confirm the target rating and construction build-up with the project team.
3. Confirm available cavity or fixing depth against the product dimensions.
4. Check NCC, fire, BAL or acoustic requirements with a qualified reviewer before specifying.
5. Record the suburb/postcode so climate-zone requirements can be checked.

## Current catalogue range

{range_rows}
## Technical data

| Property | Value | Source |
| --- | --- | --- |
| Product type | {category} | Manufacturer catalogue |
| Material | {material or "Not specified"} | Manufacturer catalogue |
| Applications | {"; ".join(applications) or "General building applications"} | Manufacturer catalogue |
| Published ratings | {rating_text} | Internal catalogue; confirm against current TDS |

## Compliance and review status

- NCC / project compliance: conditional — project-specific evidence required.
- Fire: not verified per SKU.
- BAL: not verified.
- Datasheet: {tds_url or "to be sourced"} (link audited 2026-09-05; exact product TDS may still be pending).
- SDS: {sds_url or "to be sourced"}.

## Installation overview

Use the current manufacturer instructions and the project specification. Handling, fixing and jointing details must be confirmed against the TDS for the selected SKU. {("Key limitation: " + limitations[0]) if limitations else ""}

## Safety and handling

Confirm the current SDS before handling or cutting. No product-specific hazard classification is asserted here.

## Sustainability and indoor environment

Sustainability and VOC statements are manufacturer-published claims and are not independently verified in this draft. Confirm any recycled-content or Green Star wording with the manufacturer before publication.

## Warranty, returns and support

No product-specific warranty term is asserted in this draft. Refer to the manufacturer's general terms and confirm warranty wording before publication.

## Specification starting point

> Insulation shall be {name} ({category.lower()}) by {manufacturer}, selected to suit the confirmed application and required rating. Exact product, thickness and compliance to be confirmed by the project reviewer against the current manufacturer TDS before ordering.

## Source register

- SRC-01 Manufacturer datasheet: {tds_url or "to be sourced"}.
- SRC-02 Safety data sheet: {sds_url or "to be sourced"}.
- SRC-03 Internal SKU catalogue ({len(skus) if not skus.empty else len(grade_rows)} variant(s)).

## Review actions before publication

- Confirm the exact product TDS deep link and re-validate all technical values against it.
- Confirm SDS currency and handling guidance.
- Approve environmental and warranty wording with the manufacturer.
- Human review of NCC, fire, BAL and acoustic claims remains mandatory.

_Draft generated 2026-09-05 from the internal knowledge base. Technical values and claims remain subject to manufacturer review before publication._
"""
    return md, meta


def build_docx(meta: dict, out_path: Path) -> None:
    from docx import Document
    from docx.shared import Pt

    doc = Document()
    doc.add_paragraph(meta["category"].upper() if meta["category"] else "INSULATION")
    doc.add_paragraph(meta["name"]).runs[0].font.size = Pt(22)
    doc.add_paragraph(meta["tagline"].capitalize())

    doc.add_heading("1. Product description", level=1)
    doc.add_paragraph(meta["description"] or f"{meta['name']} is a {meta['category'].lower()} product family from {meta['manufacturer']}.")
    doc.add_heading("Key features", level=2)
    for feature in meta["features"]:
        doc.add_paragraph(feature + ".", style="List Bullet")

    doc.add_heading("2. Applications and selection", level=1)
    for app in meta["applications"]:
        doc.add_paragraph(app, style="List Bullet")
    doc.add_heading("Selection checklist", level=2)
    for step in [
        "Confirm the application matches the family.",
        "Confirm the target rating and construction build-up.",
        "Confirm available depth against product dimensions.",
        "Check NCC, fire, BAL or acoustic requirements with a qualified reviewer.",
        "Record the suburb/postcode for climate-zone checks.",
    ]:
        doc.add_paragraph(step, style="List Number")

    doc.add_heading("3. Current catalogue range", level=1)
    skus = meta["skus"]
    if not skus.empty:
        table = doc.add_table(rows=1, cols=3)
        table.style = "Light Grid Accent 1"
        for cell, header in zip(table.rows[0].cells, ["SKU", "Product", "Published rating"]):
            cell.text = header
        for _, sku in skus.head(25).iterrows():
            cells = table.add_row().cells
            cells[0].text = clean(sku["our_sku"])
            cells[1].text = clean(sku["product_name"])
            cells[2].text = rating_summary(pd.Series([sku["thermal_r_value"], sku["acoustic_rw"], sku["nrc_aw"]]))
    else:
        doc.add_paragraph("Range not yet extracted; confirm variants against the current manufacturer TDS.")

    doc.add_heading("4. Technical data", level=1)
    table = doc.add_table(rows=1, cols=3)
    table.style = "Light Grid Accent 1"
    for cell, header in zip(table.rows[0].cells, ["Property", "Value", "Source"]):
        cell.text = header
    for prop, value, source in [
        ("Product type", meta["category"], "Manufacturer catalogue"),
        ("Material", meta["material"] or "Not specified", "Manufacturer catalogue"),
        ("Applications", "; ".join(meta["applications"]) or "General building applications", "Manufacturer catalogue"),
    ]:
        cells = table.add_row().cells
        cells[0].text, cells[1].text, cells[2].text = prop, value, source

    doc.add_heading("5. Fire, testing and compliance context", level=1)
    doc.add_paragraph("NCC / project compliance: conditional - project-specific evidence required. Fire: not verified per SKU. BAL: not verified.")

    doc.add_heading("6. Installation overview", level=1)
    doc.add_paragraph("Use the current manufacturer instructions and the project specification. Handling, fixing and jointing details must be confirmed against the TDS for the selected SKU.")

    doc.add_heading("7. Safety, handling and SDS status", level=1)
    doc.add_paragraph(f"SDS: {meta['sds_url'] or 'to be sourced'}. Confirm the current SDS before handling or cutting.")

    doc.add_heading("8. Sustainability and indoor environment", level=1)
    doc.add_paragraph("Sustainability and VOC statements are manufacturer-published claims and are not independently verified in this draft.")

    doc.add_heading("9. Warranty, returns and support", level=1)
    doc.add_paragraph("No product-specific warranty term is asserted in this draft. Refer to the manufacturer's general terms.")

    doc.add_heading("10. Specification starting point", level=1)
    doc.add_paragraph(f"Insulation shall be {meta['name']} ({meta['category'].lower()}) by {meta['manufacturer']}, selected to suit the confirmed application and required rating. Exact product, thickness and compliance to be confirmed by the project reviewer against the current manufacturer TDS before ordering.")

    doc.add_heading("11. Public document register", level=1)
    doc.add_paragraph(f"SRC-01 Manufacturer datasheet: {meta['tds_url'] or 'to be sourced'}.")
    doc.add_paragraph(f"SRC-02 Safety data sheet: {meta['sds_url'] or 'to be sourced'}.")

    doc.add_heading("12. Technical-review actions before publication", level=1)
    for action in [
        "Confirm the exact product TDS deep link and re-validate all technical values.",
        "Confirm SDS currency and handling guidance.",
        "Approve environmental and warranty wording with the manufacturer.",
        "Human review of NCC, fire, BAL and acoustic claims remains mandatory.",
    ]:
        doc.add_paragraph(action, style="List Bullet")

    doc.add_paragraph("Draft generated 2026-09-05 from the internal knowledge base. Technical values and claims remain subject to manufacturer review before publication.")
    doc.save(out_path)


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--only", help="limit to one manufacturer directory name")
    parser.add_argument("--dry-run", action="store_true")
    args = parser.parse_args()

    skus = pd.read_csv(SKU_CSV).fillna("")
    state = json.loads(STATE_FILE.read_text(encoding="utf-8")) if STATE_FILE.exists() else {}

    written = skipped = 0
    for path in sorted(ROOT.glob("knowledge/*/families.json")):
        manufacturer_dir = path.parent.name
        if args.only and manufacturer_dir.casefold() != args.only.casefold():
            continue
        data = json.loads(path.read_text(encoding="utf-8"))
        out_dir = OUT_DIR / manufacturer_dir.replace(" ", "_")
        used_slugs: set[str] = set()
        for family in data["families"]:
            family.setdefault("manufacturer", title_case_manufacturer(manufacturer_dir))
            md_path = path.parent / family.get("knowledge_file", "")
            if not md_path.exists():
                continue
            text = md_path.read_text(encoding="utf-8")
            family_skus = skus[skus["family_id"] == family["family_id"]]
            fingerprint = hashlib.sha256((GENERATOR_VERSION + text + json.dumps(family, sort_keys=True) + str(len(family_skus))).encode()).hexdigest()
            slug = slugify(family["name"])
            base_slug = slug
            suffix = 2
            while slug in used_slugs:
                slug = f"{base_slug}_{suffix}"
                suffix += 1
            used_slugs.add(slug)
            state_key = f"{manufacturer_dir}/{slug}"
            if state.get(state_key) == fingerprint and (out_dir / f"{slug}.md").exists():
                skipped += 1
                continue
            if args.dry_run:
                print(f"would generate {state_key}")
                written += 1
                continue
            fm = parse_front_matter(text)
            md, meta = build_markdown(family, fm, text, family_skus)
            out_dir.mkdir(parents=True, exist_ok=True)
            (out_dir / f"{slug}.md").write_text(md, encoding="utf-8")
            try:
                build_docx(meta, out_dir / f"{slug}.docx")
            except Exception as exc:  # noqa: BLE001 - docx failure must not stop the batch
                print(f"  docx failed for {slug}: {exc}")
            state[state_key] = fingerprint
            written += 1

    print(f"\ngenerated: {written}, unchanged: {skipped}")
    if not args.dry_run:
        OUT_DIR.mkdir(parents=True, exist_ok=True)
        STATE_FILE.write_text(json.dumps(state, indent=2), encoding="utf-8")


if __name__ == "__main__":
    main()
